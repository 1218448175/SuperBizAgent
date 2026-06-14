"""文档解析服务模块 — 将二进制/结构化文档转换为纯文本

支持格式: TXT, MD, PDF, DOCX, HTML, CSV, XLSX
内建错误修复机制: 主解析器失败后自动尝试回退策略（如编码检测、替代解析器）

用法:
    from app.services.document_parser_service import document_parser_service
    text = document_parser_service.parse("report.pdf")
"""

from __future__ import annotations

import csv
from dataclasses import dataclass, field
from pathlib import Path
from typing import Callable, Dict, List, Set

from loguru import logger


# ═══════════════════════════════════════════════════════════════════════════════
# 公开常量
# ═══════════════════════════════════════════════════════════════════════════════

SUPPORTED_EXTENSIONS: Set[str] = {
    ".txt", ".md", ".pdf", ".docx", ".html", ".htm", ".csv", ".xlsx",
}
SUPPORTED_EXTENSION_NAMES: Set[str] = {e.lstrip(".") for e in SUPPORTED_EXTENSIONS}


# ═══════════════════════════════════════════════════════════════════════════════
# 数据模型
# ═══════════════════════════════════════════════════════════════════════════════

@dataclass
class RepairAttempt:
    """单次修复尝试的记录"""
    strategy: str         # 修复策略名称
    success: bool = False
    error: str | None = None


# ═══════════════════════════════════════════════════════════════════════════════
# 解析服务
# ═══════════════════════════════════════════════════════════════════════════════

class DocumentParserService:
    """文档解析服务 — 根据不同文件格式提取纯文本内容

    核心设计:
    1. 字典驱动调度 — 扩展名映射到解析函数，新增格式只需 register_parser()
    2. 可选依赖 — 所有非标准库依赖均为可选，ImportError 时给出 pip install 提示
    3. 修复机制 — 主解析失败 → 诊断错误类型 → 尝试一次回退策略 → 成功/返回详细错误
    """

    def __init__(self):
        self._PARSER_REGISTRY: Dict[str, Callable[[Path], str]] = {}
        self._DEP_GROUPS: Dict[str, str] = {}
        self._register_builtin_parsers()
        logger.info(
            f"DocumentParserService 已初始化，"
            f"支持 {len(self._PARSER_REGISTRY)} 种格式: "
            f"{', '.join(sorted(self._PARSER_REGISTRY.keys()))}"
        )

    # ── 公开接口 ──────────────────────────────────────────────────

    def register_parser(
        self, extension: str, parser_fn: Callable[[Path], str], dep_group: str = ""
    ) -> None:
        """注册自定义解析器（公开扩展点）

        Args:
            extension: 文件扩展名（含点，小写），如 ".pdf"
            parser_fn: 解析函数，接收 Path 返回 str
            dep_group: 依赖组名称，用于错误提示（如 "pdf"）
        """
        ext = extension.lower() if extension.startswith(".") else f".{extension.lower()}"
        self._PARSER_REGISTRY[ext] = parser_fn
        if dep_group:
            self._DEP_GROUPS[ext] = dep_group
        logger.debug(f"注册解析器: {ext} (dep_group={dep_group or 'core'})")

    def parse(self, file_path: str) -> str:
        """将文件解析为纯文本（含修复重试）

        Args:
            file_path: 文件路径

        Returns:
            str: 提取的纯文本内容

        Raises:
            ValueError: 文件不存在或不支持的格式（调用方错误）
            RuntimeError: 解析失败（环境/数据错误），含详细修复记录
        """
        path = Path(file_path).resolve()

        # 阶段 0: 前置校验
        if not path.exists():
            raise ValueError(f"文件不存在: {file_path}")

        ext = path.suffix.lower()
        parser = self._PARSER_REGISTRY.get(ext)
        if parser is None:
            supported = ", ".join(sorted(self._PARSER_REGISTRY.keys()))
            raise ValueError(f"不支持的文件格式 '{ext}'，当前支持: {supported}")

        # ── 阶段 1: 主解析器尝试 ──
        try:
            text = parser(path)
            if text.strip():
                logger.debug(f"解析成功 [{path.name}] ({ext})")
                return text
            # 主解析器成功但返回空文本（可能是空白/扫描版 PDF）
            logger.warning(f"主解析器返回空文本 [{path.name}]，尝试修复...")
            primary_error = RuntimeError("主解析器返回空文本")
        except ImportError:
            raise  # 依赖缺失不由修复处理，直接上抛
        except Exception as e:
            primary_error = e
            logger.warning(f"主解析失败 [{path.name}] ({ext}): {primary_error}")

        # ── 阶段 2: 诊断 + 修复尝试（最多一次） ──
        repair_attempts: List[RepairAttempt] = []
        repair_plan = self._diagnose(path, primary_error)

        for plan in repair_plan:
            attempt = RepairAttempt(strategy=plan)
            try:
                text = self._execute_repair(plan, path)
                if text.strip():
                    attempt.success = True
                    logger.info(f"✅ 修复成功 [{path.name}]: {plan}")
                    return text
                attempt.error = "修复后文本为空"
            except ImportError as e:
                attempt.error = f"依赖缺失: {e}"
            except Exception as e:
                attempt.error = str(e)

            repair_attempts.append(attempt)
            logger.warning(f"❌ 修复失败 [{path.name}]: {plan} → {attempt.error}")
            break  # 只尝试一次，不无限重试

        # ── 阶段 3: 全部失败，组装详细错误 ──
        detail_lines = [f"无法解析文件 '{path.name}' ({ext})"]
        detail_lines.append(f"  原始错误: {primary_error}")
        detail_lines.append(f"  尝试修复: {len(repair_attempts)} 次，全部失败")
        for a in repair_attempts:
            detail_lines.append(f"    - {a.strategy}: {a.error}")
        # 给出安装依赖的提示
        dep_group = self._DEP_GROUPS.get(ext)
        if dep_group:
            detail_lines.append(f"  提示: pip install super-biz-agent[{dep_group}]")

        raise RuntimeError("\n".join(detail_lines))

    # ── 诊断引擎 ──────────────────────────────────────────────────

    def _diagnose(self, path: Path, error: Exception) -> List[str]:
        """根据错误类型和文件扩展名生成修复策略列表（按优先级排序）"""
        ext = path.suffix.lower()
        error_str = str(error).lower()
        plans: List[str] = []

        # 通用修复: 编码检测
        if isinstance(error, (UnicodeDecodeError, LookupError)) or "codec" in error_str:
            plans.append("encoding_detect")

        # PDF 专用修复链
        if ext == ".pdf":
            plans.append("pdf_fallback")

        # DOCX 专用修复链
        if ext == ".docx":
            plans.append("docx_fallback")

        # HTML 专用修复链
        if ext in (".html", ".htm"):
            plans.append("html_fallback")

        # CSV 专用修复链
        if ext == ".csv":
            plans.append("csv_sniff")

        # 如果没有任何专用修复，且原始错误是“空文本”，尝试编码检测
        if not plans and "空文本" in str(error):
            plans.append("encoding_detect")

        if not plans:
            logger.debug(f"无可用修复策略 [{path.name}] ({ext}): {error}")

        return plans

    def _execute_repair(self, strategy: str, path: Path) -> str:
        """执行指定的修复策略"""
        repair_fns: Dict[str, Callable[[Path], str]] = {
            "encoding_detect": self._repair_encoding,
            "pdf_fallback": self._repair_pdf_fallback,
            "docx_fallback": self._repair_docx_fallback,
            "html_fallback": self._repair_html_fallback,
            "csv_sniff": self._repair_csv_sniff,
        }
        fn = repair_fns.get(strategy)
        if fn is None:
            raise ValueError(f"未知修复策略: {strategy}")
        return fn(path)

    # ── 内置解析器注册 ──────────────────────────────────────────────

    def _register_builtin_parsers(self) -> None:
        """注册所有内置解析器"""
        # Core（零依赖）
        self.register_parser(".txt", self._parse_text, dep_group="")
        self.register_parser(".md", self._parse_text, dep_group="")
        self.register_parser(".csv", self._parse_csv, dep_group="")

        # Optional（非标准库依赖）
        self.register_parser(".pdf", self._parse_pdf, dep_group="pdf")
        self.register_parser(".docx", self._parse_docx, dep_group="docx")
        self.register_parser(".html", self._parse_html, dep_group="html")
        self.register_parser(".htm", self._parse_html, dep_group="html")
        self.register_parser(".xlsx", self._parse_xlsx, dep_group="xlsx")

    # ── 内置解析器实现 ──────────────────────────────────────────────

    @staticmethod
    def _parse_text(path: Path) -> str:
        """解析纯文本 / Markdown 文件"""
        return path.read_text(encoding="utf-8")

    @staticmethod
    def _parse_pdf(path: Path) -> str:
        """解析 PDF 文件（主解析器: pymupdf）"""
        try:
            import pymupdf
        except ImportError as e:
            raise ImportError(
                "pymupdf (AGPL-3.0) 是解析 PDF 的必需依赖. "
                "安装: pip install super-biz-agent[pdf]"
            ) from e
        doc = pymupdf.open(str(path))
        try:
            pages = [page.get_text() for page in doc]
            return "\n\n".join(p for p in pages if p.strip())
        finally:
            doc.close()

    @staticmethod
    def _parse_docx(path: Path) -> str:
        """解析 DOCX 文件（主解析器: docx2txt）"""
        try:
            import docx2txt
        except ImportError as e:
            raise ImportError(
                "docx2txt 是解析 DOCX 的必需依赖. "
                "安装: pip install super-biz-agent[docx]"
            ) from e
        return docx2txt.process(str(path)) or ""

    @staticmethod
    def _parse_html(path: Path) -> str:
        """解析 HTML 文件（主解析器: beautifulsoup4 + lxml）

        使用二进制模式打开，让 lxml 自动检测 charset。
        """
        try:
            from bs4 import BeautifulSoup
        except ImportError as e:
            raise ImportError(
                "beautifulsoup4 + lxml 是解析 HTML 的必需依赖. "
                "安装: pip install super-biz-agent[html]"
            ) from e

        with open(path, "rb") as f:
            soup = BeautifulSoup(f.read(), "lxml")
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        return soup.get_text(separator="\n", strip=True)

    @staticmethod
    def _parse_csv(path: Path) -> str:
        """解析 CSV 文件（内置 csv 模块，零依赖）"""
        rows: List[str] = []
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            for row in csv.reader(f):
                if row:
                    rows.append(" | ".join(c.strip() for c in row))
        return "\n".join(rows)

    @staticmethod
    def _parse_xlsx(path: Path) -> str:
        """解析 XLSX 文件（主解析器: openpyxl，read_only 流式模式）"""
        try:
            import openpyxl
        except ImportError as e:
            raise ImportError(
                "openpyxl 是解析 XLSX 的必需依赖. "
                "安装: pip install super-biz-agent[xlsx]"
            ) from e

        wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
        try:
            parts: List[str] = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                lines: List[str] = []
                for row in ws.iter_rows(values_only=True):
                    cells = [str(c) if c is not None else "" for c in row]
                    line = " | ".join(cells).strip()
                    if line:
                        lines.append(line)
                if lines:
                    parts.append(f"=== Sheet: {sheet_name} ===")
                    parts.extend(lines)
            return "\n".join(parts)
        finally:
            wb.close()

    # ═══════════════════════════════════════════════════════════════════════
    # 修复策略实现
    # ═══════════════════════════════════════════════════════════════════════

    @staticmethod
    def _repair_encoding(path: Path) -> str:
        """编码修复: 使用 chardet 自动检测文件编码"""
        try:
            import chardet
        except ImportError:
            raise ImportError(
                "chardet 是编码修复的必需依赖. "
                "安装: pip install chardet"
            )

        with open(path, "rb") as f:
            raw = f.read()
        detected = chardet.detect(raw)
        encoding = detected.get("encoding", "utf-8")
        confidence = detected.get("confidence", 0)
        logger.info(
            f"编码检测 [{path.name}]: {encoding} "
            f"(置信度: {confidence:.0%})"
        )
        return raw.decode(encoding, errors="replace")

    @staticmethod
    def _repair_pdf_fallback(path: Path) -> str:
        """PDF 回退: pymupdf 失败 → 使用 pypdf"""
        try:
            from pypdf import PdfReader
        except ImportError:
            raise ImportError(
                "pypdf 是 PDF 回退修复的必需依赖. "
                "安装: pip install pypdf"
            )
        reader = PdfReader(str(path))
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        return "\n\n".join(pages)

    @staticmethod
    def _repair_docx_fallback(path: Path) -> str:
        """DOCX 回退: docx2txt 失败 → 使用 python-docx"""
        try:
            from docx import Document
        except ImportError:
            raise ImportError(
                "python-docx 是 DOCX 回退修复的必需依赖. "
                "安装: pip install python-docx"
            )
        doc = Document(str(path))
        parts: List[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                parts.append("\t".join(cell.text for cell in row.cells))
        return "\n".join(parts)

    @staticmethod
    def _repair_html_fallback(path: Path) -> str:
        """HTML 回退: lxml 失败 → 使用 Python 内置 html.parser"""
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            raise ImportError(
                "beautifulsoup4 是 HTML 回退修复的必需依赖. "
                "安装: pip install beautifulsoup4"
            )
        with open(path, "rb") as f:
            soup = BeautifulSoup(f.read(), "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        return soup.get_text(separator="\n", strip=True)

    @staticmethod
    def _repair_csv_sniff(path: Path) -> str:
        """CSV 修复: 使用 csv.Sniffer 自动检测分隔符和方言"""
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            sample = f.read(8192)
        try:
            dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
        except csv.Error:
            # Sniffer 失败，尝试常见分隔符
            dialect = csv.excel
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            rows = []
            for row in csv.reader(f, dialect):
                if row:
                    rows.append(" | ".join(c.strip() for c in row))
        return "\n".join(rows)


# ═══════════════════════════════════════════════════════════════════════════════
# 全局单例
# ═══════════════════════════════════════════════════════════════════════════════

document_parser_service = DocumentParserService()
